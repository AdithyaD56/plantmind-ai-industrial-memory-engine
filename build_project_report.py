from pathlib import Path
import sys

DEPS = Path(__file__).resolve().parent / ".deps"
if DEPS.exists():
    sys.path.insert(0, str(DEPS))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "PlantMind_AI_Detailed_Report.docx"
PDF_PATH = OUT_DIR / "PlantMind_AI_Detailed_Report.pdf"
ARCH_IMG = ROOT / "PlantMind_AI_Architecture.png"


def set_font(run, name="Arial", size=11, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(p, before=0, after=6, line=1.15):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
r = title.add_run("PlantMind AI: Industrial Memory & Failure Prevention Engine")
set_font(r, name="Arial", size=22, bold=True, color="1F3A5F")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(8)
r = sub.add_run("Detailed project document aligned with Problem Statement 8")
set_font(r, name="Arial", size=11, italic=True, color="555555")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(10)
r = meta.add_run("Team Members: Dhavala V D M Adithya Naidu | D Sai Karthik | K Harsha Vardhan Reddy")
set_font(r, name="Arial", size=10, color="555555")

def heading(text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, name="Calibri", size=16 if level == 1 else 13, bold=True, color="2E74B5")
    return p


def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


heading("1. Executive Summary", 1)
para("PlantMind AI is an industrial knowledge intelligence platform designed to prevent failures before they happen. It ingests heterogeneous documents such as maintenance reports, incident logs, safety SOPs, inspection records, CSV files, and project documents, then transforms them into a continuously evolving industrial memory system.")
para("Unlike a normal document chatbot, the platform builds a knowledge graph, searches historical incidents, predicts failure risk, checks compliance, and gives cited answers to engineers. The result is an operational memory layer that helps industrial teams act early, reduce downtime, and retain lessons learned.")

heading("2. Problem Statement Alignment", 1)
para("This project is aligned with Problem Statement 8 because it addresses the challenge of extracting industrial intelligence from disconnected operational documents and turning it into actionable, proactive failure prevention.")
add_bullet(doc, "It learns from maintenance reports, incident reports, and safety documents instead of treating them as static files.")
add_bullet(doc, "It detects recurring failure patterns and surfaces risk before an unplanned shutdown occurs.")
add_bullet(doc, "It improves compliance awareness by comparing plant documents against safety and procedural requirements.")
add_bullet(doc, "It supports engineers with evidence-backed recommendations rather than generic chatbot replies.")

heading("3. Project Objective", 1)
para("The core objective of PlantMind AI is to create a living industrial memory that remembers every failure, every corrective action, and every relevant procedure.")
add_bullet(doc, "Prevent repeat failures through historical pattern matching.")
add_bullet(doc, "Make industrial knowledge searchable, structured, and connected.")
add_bullet(doc, "Provide proactive alerts that help engineers inspect at-risk assets earlier.")
add_bullet(doc, "Deliver cited assistance for maintenance, compliance, and investigation workflows.")

heading("4. System Overview", 1)
para("The platform follows a pipeline architecture: Document Upload → Document Intelligence → Knowledge Graph Builder → Failure Memory Engine → Failure Intelligence Agent → Compliance Agent → Engineer Copilot.")
add_bullet(doc, "Document Intelligence handles OCR, parsing, chunking, and structured extraction.")
add_bullet(doc, "The Knowledge Graph Builder links equipment, incidents, failures, procedures, locations, and personnel.")
add_bullet(doc, "The Failure Memory Engine stores and searches historical incidents for similarity.")
add_bullet(doc, "The Engineer Copilot answers questions using retrieved evidence and citations.")

heading("5. Architecture Diagram", 1)
para("The architecture diagram below shows the end-to-end PlantMind AI flow from industrial document ingestion to risk alerts and lessons learned.")
if ARCH_IMG.exists():
    doc.add_picture(str(ARCH_IMG), width=Inches(6.9))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last.paragraph_format.space_after = Pt(6)
else:
    para("Architecture image not found. Please place PlantMind_AI_Architecture.png in the project root.")
para("The lower layer in the diagram shows the supporting stack: Streamlit for the interface, FastAPI for the backend, ChromaDB for semantic memory, NetworkX for the graph, SQLite for persistence, and OCR/document parsing for ingestion.")

heading("6. Detailed Functional Modules", 1)
heading("6.1 Document Intelligence Agent", 2)
add_bullet(doc, "Extracts text from PDFs and text-based files.")
add_bullet(doc, "Uses OCR for scanned or image-based documents.")
add_bullet(doc, "Parses CSV and Excel files into analyzable structures.")
add_bullet(doc, "Produces structured chunks and metadata for downstream intelligence.")

heading("6.2 Knowledge Graph Builder", 2)
add_bullet(doc, "Identifies equipment IDs, locations, personnel, dates, procedures, and regulations.")
add_bullet(doc, "Connects equipment to failures and failure to root cause.")
add_bullet(doc, "Maps incident reports and SOPs into an industrial relationship graph.")
add_bullet(doc, "Creates a navigable memory structure using NetworkX.")

heading("6.3 Failure Intelligence Agent", 2)
add_bullet(doc, "Compares new evidence with historical incidents.")
add_bullet(doc, "Computes failure similarity and risk score.")
add_bullet(doc, "Finds recurring root causes and similar equipment patterns.")
add_bullet(doc, "Generates recommended actions based on the inferred risk level.")

heading("6.4 Compliance Agent", 2)
add_bullet(doc, "Checks uploaded documents against safety and SOP requirements.")
add_bullet(doc, "Highlights missing compliance elements.")
add_bullet(doc, "Generates compliance percentage and audit readiness score.")
add_bullet(doc, "Provides recommendations to close documentation gaps.")

heading("6.5 Engineer Copilot", 2)
add_bullet(doc, "Answers questions about maintenance, failures, and compliance.")
add_bullet(doc, "Uses retrieval-based evidence from the document memory layer.")
add_bullet(doc, "Returns source citations so users can validate every answer.")
add_bullet(doc, "Supports investigation and decision-making rather than open-ended chat.")

heading("7. Tech Stack", 1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.autofit = False
table.columns[0].width = Inches(2.0)
table.columns[1].width = Inches(4.5)
hdr = table.rows[0].cells
shade_cell(hdr[0], "D9E2F3")
shade_cell(hdr[1], "D9E2F3")
set_cell_text(hdr[0], "Layer", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(hdr[1], "Technology Used", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
rows = [
    ("Frontend", "Streamlit, Plotly, AgGrid"),
    ("Backend", "FastAPI"),
    ("AI / Workflow", "LangGraph, LangChain, Groq"),
    ("Vector Memory", "ChromaDB"),
    ("Knowledge Graph", "NetworkX"),
    ("Database", "SQLite"),
    ("OCR", "EasyOCR"),
    ("Document Parsing", "PyMuPDF, Unstructured"),
    ("ML", "Scikit-learn"),
]
for a, b in rows:
    c = table.add_row().cells
    set_cell_text(c[0], a, bold=True, size=10)
    set_cell_text(c[1], b, size=10)

heading("8. Sample Demo Flow", 1)
para("The prototype is designed for a live demonstration where a judge uploads industrial documents and immediately sees the memory engine in action.")
add_number(doc, "Upload a maintenance report, an incident report, and a safety SOP.")
add_number(doc, "Show that OCR, extraction, and graph creation happen automatically.")
add_number(doc, "Open the Failure Intelligence view for equipment such as Pump P101.")
add_number(doc, "Ask the Engineer Copilot why the asset is at risk and show the citations.")
add_number(doc, "Finish by highlighting the proactive warning and recommended inspection window.")

heading("9. Project Impact", 1)
add_bullet(doc, "Reduces repeated failures by learning from historical industrial events.")
add_bullet(doc, "Improves uptime by turning documents into actionable maintenance intelligence.")
add_bullet(doc, "Supports safer operations through compliance and procedure awareness.")
add_bullet(doc, "Preserves organizational memory even when staff and shift teams change.")

heading("10. Conclusion", 1)
para("PlantMind AI is built as an industrial operating system for knowledge. It combines document intelligence, knowledge graphs, failure memory, compliance checking, and an evidence-based copilot to help industries prevent failures before they happen.")
para("This makes the system highly aligned with Problem Statement 8 because it does not merely retrieve documents; it learns from them, connects them, and converts them into proactive operational intelligence.")

heading("Team Members", 1)
add_bullet(doc, "Dhavala V D M Adithya Naidu")
add_bullet(doc, "D Sai Karthik")
add_bullet(doc, "K Harsha Vardhan Reddy")

doc.save(DOCX_PATH)
print(f"Wrote {DOCX_PATH}")
